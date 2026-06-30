"""Extended tests for jobhunt/resume_renderer.py — covering the helpers,
text_to_* functions, and the DOCXRenderer/PDFRenderer that were at 44% coverage.
"""

from __future__ import annotations

import pytest

from jobhunt.resume_renderer import (
    HTMLRenderer,
    PDFRenderer,
    RendererUnavailable,
    TextRenderer,
    _draft_to_html,
    _esc,
    _is_bullet,
    _is_heading,
    _bullet_text,
    _latin1,
    _parse_runs,
    get_renderer,
    text_to_pdf,
    text_to_styled_html,
)
from jobhunt.resume_template import Bullet, ResumeDraft, ResumeSection


# ----------------------------------------------------------------- fixtures


def _simple_draft() -> ResumeDraft:
    return ResumeDraft(
        candidate_name="Jane Doe",
        candidate_email="jane@example.com",
        target_role="Backend Engineer",
        target_company="Acme",
        summary="Experienced backend engineer with 5 years in Python.",
        sections=[
            ResumeSection(
                title="Experience",
                kind="experience",
                bullets=[Bullet(text="Built microservices", evidence_id="e1")],
                body="",
                rows=[
                    {
                        "left": "**Senior Engineer, Globex**",
                        "right": "2020 - Present",
                        "bullets": [
                            {"text": "Led migration to Kubernetes", "evidence_id": "e2"}
                        ],
                    }
                ],
            ),
            ResumeSection(
                title="Skills",
                kind="skills",
                body="Python, Go, Kubernetes, PostgreSQL",
                bullets=[],
                rows=[],
            ),
            ResumeSection(
                title="Education",
                kind="education",
                bullets=[],
                body="",
                rows=[{"left": "**MIT** — BSc CS", "right": "2016", "link": "", "bullets": []}],
            ),
        ],
        matched_keywords=["python", "kubernetes"],
        missing_keywords=["rust"],
        phone="555-0123",
        location="Remote",
        links={"github": "github.com/janedoe"},
    )


# ----------------------------------------------------------------- helpers


class TestHelpers:
    def test_latin1_transliterates_unicode(self):
        assert _latin1("hello \u2014 world") == "hello - world"
        assert _latin1("\u2018quoted\u2019") == "'quoted'"
        assert _latin1("\u201csmart\u201d") == '"smart"'
        assert _latin1("bullet \u2022 point") == "bullet - point"
        assert _latin1("arrow \u2192 left") == "arrow -> left"
        assert _latin1("\u2265 5 and \u2264 10") == ">= 5 and <= 10"

    def test_latin1_handles_plain_ascii(self):
        assert _latin1("plain text") == "plain text"

    def test_latin1_replaces_unknown_unicode(self):
        # Characters outside latin-1 replaced with ?
        result = _latin1("こんにちは")
        assert "?" in result

    def test_esc_html_entities(self):
        assert _esc("a & b") == "a &amp; b"
        assert _esc("<script>") == "&lt;script&gt;"
        assert _esc("normal") == "normal"

    def test_is_heading_all_caps_short(self):
        assert _is_heading("EXPERIENCE")
        assert _is_heading("SKILLS")
        assert _is_heading("  EDUCATION  ")

    def test_is_heading_rejects_long_text(self):
        assert not _is_heading("A" * 41)

    def test_is_heading_rejects_lowercase(self):
        assert not _is_heading("Experience")
        assert not _is_heading("mixed CASE text")

    def test_is_heading_rejects_empty(self):
        assert not _is_heading("")
        assert not _is_heading("   ")

    def test_is_heading_rejects_numbers_only(self):
        assert not _is_heading("12345")

    def test_is_bullet_various_markers(self):
        assert _is_bullet("- item")
        assert _is_bullet("* item")
        assert _is_bullet("• item")
        assert _is_bullet("  - indented")

    def test_is_bullet_rejects_non_bullets(self):
        assert not _is_bullet("normal line")
        assert not _is_bullet("")
        assert not _is_bullet("--flag")

    def test_bullet_text_strips_marker(self):
        assert _bullet_text("- built stuff") == "built stuff"
        assert _bullet_text("* item here") == "item here"
        assert _bullet_text("  - indented") == "indented"

    def test_parse_runs_plain_text(self):
        runs = _parse_runs("hello world")
        assert runs == [("hello world", False)]

    def test_parse_runs_bold(self):
        runs = _parse_runs("**bold** text")
        assert ("bold", True) in runs
        assert ("text", False) in [(t.strip(), b) for t, b in runs if t.strip()]

    def test_parse_runs_multiple_bold(self):
        runs = _parse_runs("**a** and **b**")
        bold_texts = [t for t, b in runs if b]
        assert "a" in bold_texts
        assert "b" in bold_texts

    def test_parse_runs_empty_string(self):
        runs = _parse_runs("")
        assert runs == [("", False)]


# ----------------------------------------------------------------- text_to_pdf


class TestTextToPdf:
    def test_basic_rendering(self):
        try:
            result = text_to_pdf("Jane Doe", "EXPERIENCE\n- Built services\n- Led team")
        except RendererUnavailable:
            pytest.skip("fpdf2 not installed")
        assert result[:4] == b"%PDF"
        assert len(result) > 200

    def test_with_headings_and_bullets(self):
        try:
            body = "EXPERIENCE\n- item 1\n- item 2\n\nSKILLS\nPython, Go\n\nEDUCATION\nMIT BSc"
            result = text_to_pdf("Test Person", body)
        except RendererUnavailable:
            pytest.skip("fpdf2 not installed")
        assert result[:4] == b"%PDF"

    def test_empty_body(self):
        try:
            result = text_to_pdf("Name", "")
        except RendererUnavailable:
            pytest.skip("fpdf2 not installed")
        assert result[:4] == b"%PDF"

    def test_unicode_transliteration(self):
        try:
            result = text_to_pdf("José García", "EXPERIENCE\n- Built → deployed")
        except RendererUnavailable:
            pytest.skip("fpdf2 not installed")
        assert result[:4] == b"%PDF"

    def test_blank_lines_handled(self):
        try:
            body = "EXPERIENCE\n\n- item\n\n\nSKILLS\nPython"
            result = text_to_pdf("Name", body)
        except RendererUnavailable:
            pytest.skip("fpdf2 not installed")
        assert result[:4] == b"%PDF"


# ----------------------------------------------------------------- text_to_styled_html


class TestTextToStyledHtml:
    def test_basic_rendering(self):
        html = text_to_styled_html("Jane Doe", "EXPERIENCE\n- Built services\n\nSKILLS\nPython")
        assert "Jane Doe" in html
        assert "<h2>" in html
        assert "<li>" in html
        assert "Python" in html

    def test_headings_are_titlecased(self):
        html = text_to_styled_html("Name", "EXPERIENCE\n- stuff")
        assert "Experience" in html

    def test_bullets_in_ul(self):
        html = text_to_styled_html("Name", "SECTION\n- bullet one\n- bullet two")
        assert "<ul>" in html
        assert "bullet one" in html
        assert "bullet two" in html

    def test_paragraph_text(self):
        html = text_to_styled_html("Name", "SUMMARY\nSome paragraph text.")
        assert "<p>" in html
        assert "Some paragraph text." in html

    def test_tab_title_override(self):
        html = text_to_styled_html("Name", "body", tab_title="Custom Title")
        assert "<title>Custom Title</title>" in html

    def test_default_tab_title_uses_heading(self):
        html = text_to_styled_html("My Name", "body")
        assert "<title>My Name</title>" in html

    def test_html_escaping(self):
        html = text_to_styled_html("A & B", "SECTION\n- <script>alert(1)</script>")
        assert "&amp;" in html
        assert "&lt;script&gt;" in html

    def test_empty_body(self):
        html = text_to_styled_html("Name", "")
        assert "Name" in html
        assert "<!doctype html>" in html


# ----------------------------------------------------------------- text_to_docx


class TestTextToDocx:
    def test_basic_rendering(self):
        from jobhunt.resume_renderer import text_to_docx

        try:
            result = text_to_docx("Jane Doe", "EXPERIENCE\n- Built services\n\nSKILLS\nPython")
        except RendererUnavailable:
            pytest.skip("python-docx not installed")
        # DOCX files start with PK (zip format)
        assert result[:2] == b"PK"
        assert len(result) > 100

    def test_empty_body(self):
        from jobhunt.resume_renderer import text_to_docx

        try:
            result = text_to_docx("Name", "")
        except RendererUnavailable:
            pytest.skip("python-docx not installed")
        assert result[:2] == b"PK"

    def test_sections_and_bullets(self):
        from jobhunt.resume_renderer import text_to_docx

        try:
            body = "EXPERIENCE\n- item 1\n- item 2\n\nSKILLS\nPython, Go"
            result = text_to_docx("Person", body)
        except RendererUnavailable:
            pytest.skip("python-docx not installed")
        assert result[:2] == b"PK"


# ----------------------------------------------------------------- DOCXRenderer


class TestDOCXRenderer:
    def test_render_basic(self):
        from jobhunt.resume_renderer import DOCXRenderer

        draft = _simple_draft()
        renderer = DOCXRenderer()
        try:
            result = renderer.render(draft)
        except RendererUnavailable:
            pytest.skip("python-docx not installed")
        assert result[:2] == b"PK"
        assert len(result) > 100

    def test_renderer_attributes(self):
        from jobhunt.resume_renderer import DOCXRenderer

        r = DOCXRenderer()
        assert r.extension == "docx"
        assert "wordprocessingml" in r.content_type


# ----------------------------------------------------------------- PDFRenderer


class TestPDFRenderer:
    def test_render_basic(self):
        draft = _simple_draft()
        renderer = PDFRenderer()
        try:
            result = renderer.render(draft)
        except RendererUnavailable:
            pytest.skip("fpdf2 not installed")
        assert result[:4] == b"%PDF"

    def test_renderer_attributes(self):
        r = PDFRenderer()
        assert r.extension == "pdf"
        assert r.content_type == "application/pdf"


# ----------------------------------------------------------------- HTMLRenderer


class TestHTMLRenderer:
    def test_render_basic(self):
        draft = _simple_draft()
        renderer = HTMLRenderer()
        result = renderer.render(draft)
        html = result.decode("utf-8")
        assert "Jane Doe" in html
        assert "Backend Engineer" in html
        assert "<html>" in html

    def test_renderer_attributes(self):
        r = HTMLRenderer()
        assert r.extension == "html"
        assert r.content_type == "text/html"


# ----------------------------------------------------------------- TextRenderer


class TestTextRenderer:
    def test_render_basic(self):
        draft = _simple_draft()
        renderer = TextRenderer()
        result = renderer.render(draft)
        text = result.decode("utf-8")
        assert "Jane Doe" in text

    def test_renderer_attributes(self):
        r = TextRenderer()
        assert r.extension == "txt"
        assert r.content_type == "text/plain"


# ----------------------------------------------------------------- _draft_to_html


class TestDraftToHtml:
    def test_includes_metadata(self):
        draft = _simple_draft()
        html = _draft_to_html(draft)
        assert "Jane Doe" in html
        assert "jane@example.com" in html
        assert "Backend Engineer" in html
        assert "Acme" in html

    def test_includes_summary(self):
        draft = _simple_draft()
        html = _draft_to_html(draft)
        assert "Experienced backend engineer" in html

    def test_includes_sections_and_bullets(self):
        draft = _simple_draft()
        html = _draft_to_html(draft)
        assert "<h2>Experience</h2>" in html
        assert "Built microservices" in html
        assert 'data-evidence-id="e1"' in html

    def test_section_body(self):
        draft = _simple_draft()
        html = _draft_to_html(draft)
        assert "Python, Go, Kubernetes" in html


# ----------------------------------------------------------------- get_renderer


class TestGetRenderer:
    def test_known_formats(self):
        assert isinstance(get_renderer("txt"), TextRenderer)
        assert isinstance(get_renderer("html"), HTMLRenderer)
        assert isinstance(get_renderer("pdf"), PDFRenderer)

    def test_unknown_format_falls_back_to_text(self):
        r = get_renderer("unknown")
        assert isinstance(r, TextRenderer)

    def test_docx_format(self):
        from jobhunt.resume_renderer import DOCXRenderer

        assert isinstance(get_renderer("docx"), DOCXRenderer)


# ----------------------------------------------------------------- draft_to_styled_html


class TestDraftToStyledHtml:
    def test_renders_name_and_contact(self):
        from jobhunt.resume_renderer import draft_to_styled_html

        draft = _simple_draft()
        html = draft_to_styled_html(draft)
        assert "Jane Doe" in html
        assert "jane@example.com" in html
        assert "555-0123" in html

    def test_renders_sections(self):
        from jobhunt.resume_renderer import draft_to_styled_html

        draft = _simple_draft()
        html = draft_to_styled_html(draft)
        assert "Experience" in html
        assert "Skills" in html
        assert "Education" in html

    def test_skills_section_has_class(self):
        from jobhunt.resume_renderer import draft_to_styled_html

        draft = _simple_draft()
        html = draft_to_styled_html(draft)
        assert 'class="skills"' in html

    def test_rows_have_left_right(self):
        from jobhunt.resume_renderer import draft_to_styled_html

        draft = _simple_draft()
        html = draft_to_styled_html(draft)
        assert 'class="left"' in html
        assert 'class="right"' in html
        assert "2020 - Present" in html

    def test_bold_in_rows(self):
        from jobhunt.resume_renderer import draft_to_styled_html

        draft = _simple_draft()
        html = draft_to_styled_html(draft)
        assert "<strong>" in html

    def test_footer_rendered(self):
        from jobhunt.resume_renderer import draft_to_styled_html

        draft = _simple_draft()
        html = draft_to_styled_html(draft, footer="Built with JobHunt")
        assert "Built with JobHunt" in html
        assert 'class="footer"' in html

    def test_no_footer_when_none(self):
        from jobhunt.resume_renderer import draft_to_styled_html

        draft = _simple_draft()
        html = draft_to_styled_html(draft)
        assert 'class="footer"' not in html

    def test_summary_skipped_kind(self):
        from jobhunt.resume_renderer import draft_to_styled_html

        draft = _simple_draft()
        # Add a summary-kind section that should be skipped
        draft.sections.insert(0, ResumeSection(title="Summary", kind="summary", body="skip me"))
        html = draft_to_styled_html(draft)
        # The top-level summary is still rendered as .summary paragraph
        assert 'class="summary"' in html

    def test_row_with_link(self):
        from jobhunt.resume_renderer import draft_to_styled_html

        draft = _simple_draft()
        draft.sections[2].rows[0]["link"] = "https://mit.edu"
        html = draft_to_styled_html(draft)
        assert "https://mit.edu" in html
        assert "<a href=" in html


# ----------------------------------------------------------------- draft_to_pdf


class TestDraftToPdf:
    def test_renders_valid_pdf(self):
        from jobhunt.resume_renderer import draft_to_pdf

        draft = _simple_draft()
        try:
            result = draft_to_pdf(draft)
        except RendererUnavailable:
            pytest.skip("fpdf2 not installed")
        assert result[:4] == b"%PDF"
        assert len(result) > 500

    def test_renders_with_no_summary(self):
        from jobhunt.resume_renderer import draft_to_pdf

        draft = _simple_draft()
        draft.summary = ""
        try:
            result = draft_to_pdf(draft)
        except RendererUnavailable:
            pytest.skip("fpdf2 not installed")
        assert result[:4] == b"%PDF"

    def test_renders_with_skills_section(self):
        from jobhunt.resume_renderer import draft_to_pdf

        draft = _simple_draft()
        try:
            result = draft_to_pdf(draft)
        except RendererUnavailable:
            pytest.skip("fpdf2 not installed")
        # Skills section has body, not rows
        assert result[:4] == b"%PDF"

    def test_renders_with_body_in_section(self):
        from jobhunt.resume_renderer import draft_to_pdf

        draft = _simple_draft()
        draft.sections[0].body = "Some additional body text"
        try:
            result = draft_to_pdf(draft)
        except RendererUnavailable:
            pytest.skip("fpdf2 not installed")
        assert result[:4] == b"%PDF"

    def test_no_contact_line(self):
        from jobhunt.resume_renderer import draft_to_pdf

        draft = _simple_draft()
        draft.candidate_email = ""
        draft.phone = ""
        draft.location = ""
        draft.links = {}
        try:
            result = draft_to_pdf(draft)
        except RendererUnavailable:
            pytest.skip("fpdf2 not installed")
        assert result[:4] == b"%PDF"

    def test_row_without_right(self):
        from jobhunt.resume_renderer import draft_to_pdf

        draft = _simple_draft()
        draft.sections[0].rows[0].pop("right", None)
        try:
            result = draft_to_pdf(draft)
        except RendererUnavailable:
            pytest.skip("fpdf2 not installed")
        assert result[:4] == b"%PDF"

    def test_row_with_link_instead_of_right(self):
        from jobhunt.resume_renderer import draft_to_pdf

        draft = _simple_draft()
        draft.sections[2].rows[0]["link"] = "https://mit.edu"
        draft.sections[2].rows[0]["right"] = ""
        try:
            result = draft_to_pdf(draft)
        except RendererUnavailable:
            pytest.skip("fpdf2 not installed")
        assert result[:4] == b"%PDF"

    def test_bullet_as_string(self):
        from jobhunt.resume_renderer import draft_to_pdf

        draft = _simple_draft()
        # Some bullets might be plain strings instead of dicts
        draft.sections[0].rows[0]["bullets"] = ["plain string bullet"]
        try:
            result = draft_to_pdf(draft)
        except RendererUnavailable:
            pytest.skip("fpdf2 not installed")
        assert result[:4] == b"%PDF"


# ----------------------------------------------------------------- draft_to_docx


class TestDraftToDocx:
    def test_renders_valid_docx(self):
        from jobhunt.resume_renderer import draft_to_docx

        draft = _simple_draft()
        try:
            result = draft_to_docx(draft)
        except RendererUnavailable:
            pytest.skip("python-docx not installed")
        assert result[:2] == b"PK"

    def test_has_name_in_content(self):
        from jobhunt.resume_renderer import draft_to_docx

        draft = _simple_draft()
        try:
            result = draft_to_docx(draft)
        except RendererUnavailable:
            pytest.skip("python-docx not installed")
        from io import BytesIO

        from docx import Document  # type: ignore

        doc = Document(BytesIO(result))
        full = "\n".join(p.text for p in doc.paragraphs)
        assert "Jane Doe" in full

    def test_skills_body_rendered(self):
        from jobhunt.resume_renderer import draft_to_docx

        draft = _simple_draft()
        try:
            result = draft_to_docx(draft)
        except RendererUnavailable:
            pytest.skip("python-docx not installed")
        from io import BytesIO

        from docx import Document  # type: ignore

        doc = Document(BytesIO(result))
        full = "\n".join(p.text for p in doc.paragraphs)
        assert "Python, Go, Kubernetes" in full

    def test_rows_with_bold(self):
        from jobhunt.resume_renderer import draft_to_docx

        draft = _simple_draft()
        try:
            result = draft_to_docx(draft)
        except RendererUnavailable:
            pytest.skip("python-docx not installed")
        from io import BytesIO

        from docx import Document  # type: ignore

        doc = Document(BytesIO(result))
        # Check that bold runs exist
        has_bold = any(
            run.bold for p in doc.paragraphs for run in p.runs if run.bold
        )
        assert has_bold

    def test_no_contact_line(self):
        from jobhunt.resume_renderer import draft_to_docx

        draft = _simple_draft()
        draft.candidate_email = ""
        draft.phone = ""
        draft.location = ""
        draft.links = {}
        try:
            result = draft_to_docx(draft)
        except RendererUnavailable:
            pytest.skip("python-docx not installed")
        assert result[:2] == b"PK"

    def test_summary_skipped_section(self):
        from jobhunt.resume_renderer import draft_to_docx

        draft = _simple_draft()
        draft.sections.insert(0, ResumeSection(title="Summary", kind="summary", body="skip me"))
        try:
            result = draft_to_docx(draft)
        except RendererUnavailable:
            pytest.skip("python-docx not installed")
        assert result[:2] == b"PK"

    def test_bullet_as_plain_string(self):
        from jobhunt.resume_renderer import draft_to_docx

        draft = _simple_draft()
        draft.sections[0].rows[0]["bullets"] = ["plain string"]
        try:
            result = draft_to_docx(draft)
        except RendererUnavailable:
            pytest.skip("python-docx not installed")
        from io import BytesIO

        from docx import Document  # type: ignore

        doc = Document(BytesIO(result))
        full = "\n".join(p.text for p in doc.paragraphs)
        assert "plain string" in full
