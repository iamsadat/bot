"""Tests for résumé file upload-parse + GitHub project import."""

from __future__ import annotations

import base64
from io import BytesIO

import pytest

pytest.importorskip("fastapi")

from fastapi.testclient import TestClient  # noqa: E402

from jobhunt.dashboard.server import DashboardState, create_app  # noqa: E402
from jobhunt.http import FakeHTTPClient  # noqa: E402
from jobhunt.integrations.github import GitHubClient, repos_to_projects  # noqa: E402
from jobhunt.onboarding import ResumeFileError, extract_resume_text  # noqa: E402
from jobhunt.trace import ThoughtBus, TraceStore  # noqa: E402


def _client():
    st = DashboardState(trace_store=TraceStore(), bus=ThoughtBus())
    c = TestClient(create_app(st))
    c.post("/api/onboarding/profile", json={
        "name": "Ada", "email": "ada@x.com",
        "target_roles": ["backend"], "locations": ["Remote"]})
    return st, c


# ----- file extraction -----------------------------------------------------

def test_extract_txt():
    assert "hello" in extract_resume_text("cv.txt", b"hello world")


def test_extract_docx_roundtrip():
    pytest.importorskip("docx")  # python-docx is an optional dep (skip in CI)
    from docx import Document
    doc = Document()
    doc.add_paragraph("Senior Backend Engineer")
    doc.add_paragraph("Python, Kubernetes")
    buf = BytesIO()
    doc.save(buf)
    text = extract_resume_text("cv.docx", buf.getvalue())
    assert "Backend Engineer" in text and "Kubernetes" in text


def test_extract_unsupported_type_raises():
    with pytest.raises(ResumeFileError):
        extract_resume_text("cv.rtf", b"x")


# ----- GitHub client + mapping --------------------------------------------

def test_github_client_fetches_repos():
    http = FakeHTTPClient({
        "https://api.github.com/users/ada/repos?sort=updated&per_page=100":
            [{"name": "jobhunt", "fork": False}],
    })
    repos = GitHubClient(http=http).fetch_repos("@ada")  # @ stripped
    assert repos[0]["name"] == "jobhunt"


def test_repos_to_projects_skips_forks_and_maps_skills():
    repos = [
        {"name": "real", "fork": False, "language": "Python",
         "topics": ["fastapi"], "description": "A tool", "html_url": "u",
         "stargazers_count": 5},
        {"name": "forked", "fork": True, "language": "Go"},
    ]
    projs = repos_to_projects(repos)
    assert [p["name"] for p in projs] == ["real"]
    assert "python" in projs[0]["skills"] and "fastapi" in projs[0]["skills"]
    assert projs[0]["link"] == "u"


# ----- endpoints -----------------------------------------------------------

def test_parse_resume_file_endpoint_fills_skills():
    st, c = _client()
    payload = base64.b64encode(b"Skills: Python, Kubernetes, Redis").decode()
    r = c.post("/api/profile/parse-resume-file",
               json={"filename": "cv.txt", "content_base64": payload})
    assert r.status_code == 200
    assert "python" in r.json()["skills"]
    assert "python" in st.user_profile.skills


def test_parse_resume_file_rejects_unknown_type():
    _, c = _client()
    r = c.post("/api/profile/parse-resume-file",
               json={"filename": "cv.rtf", "content_base64": base64.b64encode(b"x").decode()})
    assert r.status_code == 415


def test_import_github_endpoint(monkeypatch):
    st, c = _client()
    monkeypatch.setattr(
        "jobhunt.integrations.github.GitHubClient.fetch_repos",
        lambda self, u: [{"name": "jobhunt", "fork": False, "language": "Python",
                          "html_url": "https://github.com/ada/jobhunt",
                          "description": "x", "stargazers_count": 1}],
    )
    r = c.post("/api/profile/import-github", json={"username": "ada"})
    assert r.status_code == 200 and r.json()["added"] == 1
    assert any(p["name"] == "jobhunt" for p in st.user_profile.projects)
    assert st.user_profile.links["github"].endswith("ada")
