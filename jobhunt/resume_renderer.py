"""Resume rendering: plain text (stdlib), PDF (WeasyPrint), DOCX (python-docx).

Each renderer returns ``bytes`` so the result can be uploaded straight to S3.
Optional dependencies degrade gracefully — if WeasyPrint or python-docx are
missing, the renderer raises ``RendererUnavailable`` and the caller can fall
back to text. This means tests stay dependency-free.
"""

from __future__ import annotations

from io import BytesIO
from typing import Protocol

from jobhunt.resume_template import ResumeDraft


class RendererUnavailable(RuntimeError):
    """Raised when an optional renderer dependency is missing."""


class Renderer(Protocol):
    extension: str
    content_type: str

    def render(self, draft: ResumeDraft) -> bytes: ...


# ---------------------------------------------------------- plain text


class TextRenderer:
    extension = "txt"
    content_type = "text/plain"

    def render(self, draft: ResumeDraft) -> bytes:
        return draft.to_text().encode("utf-8")


# ----------------------------------------------------------------- HTML


def _draft_to_html(draft: ResumeDraft) -> str:
    """Inline-styled HTML — works for both browser preview and WeasyPrint PDF."""
    sections_html: list[str] = []
    for sec in draft.sections:
        sec_html = [f"<h2>{sec.title}</h2>"]
        if sec.body:
            sec_html.append(f"<p>{sec.body}</p>")
        if sec.bullets:
            sec_html.append("<ul>")
            for b in sec.bullets:
                sec_html.append(
                    f'<li data-evidence-id="{b.evidence_id}">{b.text}</li>'
                )
            sec_html.append("</ul>")
        sections_html.append("\n".join(sec_html))

    return f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<title>{draft.candidate_name} — {draft.target_role}</title>
<style>
  body {{ font-family: -apple-system, Helvetica, sans-serif; max-width: 720px;
          margin: 32px auto; color: #1a1a1a; line-height: 1.45; }}
  h1 {{ font-size: 1.6em; margin-bottom: 0.2em; }}
  h2 {{ font-size: 1.1em; border-bottom: 1px solid #ccc; margin-top: 1.4em; }}
  .meta {{ color: #666; }}
  ul {{ padding-left: 1.2em; }}
  li {{ margin-bottom: 0.3em; }}
</style></head><body>
  <h1>{draft.candidate_name}</h1>
  <div class="meta">{draft.candidate_email}</div>
  <div class="meta">Target: {draft.target_role} @ {draft.target_company}</div>
  <h2>Summary</h2>
  <p>{draft.summary}</p>
  {chr(10).join(sections_html)}
</body></html>"""


class HTMLRenderer:
    extension = "html"
    content_type = "text/html"

    def render(self, draft: ResumeDraft) -> bytes:
        return _draft_to_html(draft).encode("utf-8")


# -------------------------------------------------------------- PDF (WeasyPrint)


class PDFRenderer:
    extension = "pdf"
    content_type = "application/pdf"

    def render(self, draft: ResumeDraft) -> bytes:
        # Pure-Python fpdf2 — no system libraries, installs cleanly on
        # Windows and in Docker (unlike WeasyPrint's GTK/Pango stack).
        body = draft.to_text()
        lines = body.split("\n")
        heading = lines[0].strip() if lines and lines[0].strip() else draft.candidate_name
        return text_to_pdf(heading, "\n".join(lines[1:]))


# --------------------------------------------------------------- DOCX


class DOCXRenderer:
    extension = "docx"
    content_type = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    def render(self, draft: ResumeDraft) -> bytes:
        try:
            from docx import Document  # type: ignore
        except ImportError as e:
            raise RendererUnavailable(
                "python-docx not installed; pip install python-docx"
            ) from e

        doc = Document()
        doc.add_heading(draft.candidate_name, level=0)
        doc.add_paragraph(draft.candidate_email)
        doc.add_paragraph(
            f"Target: {draft.target_role} @ {draft.target_company}"
        )

        doc.add_heading("Summary", level=2)
        doc.add_paragraph(draft.summary)

        for sec in draft.sections:
            doc.add_heading(sec.title, level=2)
            if sec.body:
                doc.add_paragraph(sec.body)
            for b in sec.bullets:
                doc.add_paragraph(b.text, style="List Bullet")

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()


# ----------------------------------------------------------------- helpers
#
# These take a heading + a structured plain-text body (the format produced by
# ResumeArchitectAgent._render_resume: ALL-CAPS section headers, "- " bullets,
# blank-line separated paragraphs) and render clean PDF / HTML / DOCX. They
# back the dashboard download endpoint so every format shares one layout.

# Core PDF fonts are latin-1 only; transliterate the handful of unicode
# punctuation our text can contain so fpdf2 never raises on an exotic glyph.
_TRANSLIT = {
    "—": "-", "–": "-", "‘": "'", "’": "'",
    "“": '"', "”": '"', "•": "-", "…": "...",
    "→": "->", "←": "<-", "≥": ">=", "≤": "<=",
    " ": " ", " ": " ", "​": "",
}


def _latin1(s: str) -> str:
    for k, v in _TRANSLIT.items():
        s = s.replace(k, v)
    return s.encode("latin-1", "replace").decode("latin-1")


def _esc(s: str) -> str:
    return (
        s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    )


def _is_heading(line: str) -> bool:
    """Section headers are short, all-caps, alphabetic lines."""
    s = line.strip()
    return bool(s) and len(s) <= 40 and s == s.upper() and any(c.isalpha() for c in s)


def _is_bullet(line: str) -> bool:
    return line.lstrip().startswith(("- ", "* ", "• "))


def _bullet_text(line: str) -> str:
    return line.lstrip()[2:].strip()


def text_to_pdf(heading: str, body: str) -> bytes:
    """Render a heading + structured body to a clean A4 PDF via fpdf2."""
    try:
        from fpdf import FPDF  # type: ignore
    except ImportError as e:
        raise RendererUnavailable("fpdf2 not installed; pip install fpdf2") from e

    pdf = FPDF(format="A4")
    pdf.set_auto_page_break(auto=True, margin=16)
    pdf.set_margins(18, 16, 18)
    pdf.add_page()

    # multi_cell leaves the cursor at the right margin by default, which makes
    # the *next* full-width cell see zero available width; LMARGIN/NEXT returns
    # the cursor to the left margin on the following line.
    def cell(text: str, h: float) -> None:
        pdf.multi_cell(0, h, _latin1(text), new_x="LMARGIN", new_y="NEXT")

    pdf.set_text_color(17, 17, 24)
    pdf.set_font("Helvetica", "B", 19)
    cell(heading, 9)
    pdf.ln(1)

    for raw in body.split("\n"):
        line = raw.rstrip()
        if not line.strip():
            pdf.ln(2)
            continue
        if _is_heading(line):
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", 11)
            pdf.set_text_color(90, 92, 120)
            cell(line.strip(), 6)
            pdf.set_text_color(30, 30, 38)
            continue
        if _is_bullet(line):
            pdf.set_font("Helvetica", size=10)
            cell("-  " + _bullet_text(line), 5)
            continue
        pdf.set_font("Helvetica", size=10)
        cell(line, 5)

    return bytes(pdf.output())


def text_to_styled_html(heading: str, body: str, *, tab_title: str = "") -> str:
    """Render a heading + structured body to a self-contained styled HTML page."""
    blocks: list[str] = []
    bullets: list[str] = []

    def flush() -> None:
        if bullets:
            blocks.append("<ul>" + "".join(f"<li>{_esc(b)}</li>" for b in bullets) + "</ul>")
            bullets.clear()

    for raw in body.split("\n"):
        line = raw.rstrip()
        if not line.strip():
            flush()
            continue
        if _is_heading(line):
            flush()
            blocks.append(f"<h2>{_esc(line.strip().title())}</h2>")
            continue
        if _is_bullet(line):
            bullets.append(_bullet_text(line))
            continue
        flush()
        blocks.append(f"<p>{_esc(line)}</p>")
    flush()

    inner = "\n  ".join(blocks)
    return f"""<!doctype html>
<html><head><meta charset="utf-8">
<title>{_esc(tab_title or heading)}</title>
<style>
  body {{ font-family: -apple-system, "Segoe UI", Helvetica, Arial, sans-serif;
          max-width: 720px; margin: 40px auto; padding: 0 24px;
          color: #14161f; line-height: 1.5; }}
  h1 {{ font-size: 1.7em; margin: 0 0 4px; letter-spacing: -0.02em; }}
  h2 {{ font-size: 0.82em; text-transform: uppercase; letter-spacing: 0.09em;
        color: #5b6478; border-bottom: 1px solid #e3e6ef;
        padding-bottom: 4px; margin: 26px 0 10px; }}
  p {{ margin: 0 0 8px; }}
  ul {{ margin: 0 0 8px; padding-left: 20px; }}
  li {{ margin-bottom: 5px; }}
</style></head><body>
  <h1>{_esc(heading)}</h1>
  {inner}
</body></html>"""


def text_to_docx(heading: str, body: str) -> bytes:
    """Render a heading + structured body to a .docx via python-docx."""
    try:
        from docx import Document  # type: ignore
    except ImportError as e:
        raise RendererUnavailable(
            "python-docx not installed; pip install python-docx"
        ) from e

    doc = Document()
    doc.add_heading(heading, level=0)
    for raw in body.split("\n"):
        line = raw.rstrip()
        if not line.strip():
            continue
        if _is_heading(line):
            doc.add_heading(line.strip().title(), level=2)
        elif _is_bullet(line):
            doc.add_paragraph(_bullet_text(line), style="List Bullet")
        else:
            doc.add_paragraph(line)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------- draft-aware layout
#
# These consume a ResumeDraft's structured ``rows`` directly (no text parsing)
# to produce the clean single-column template: centered name + contact, bold
# section headers with rules, left title/company with right-aligned dates,
# bullet lists, right-aligned project links, comma-list skills, education lines.


def _parse_runs(s: str) -> list[tuple[str, bool]]:
    """Split a ``**bold**``-marked string into (text, is_bold) runs."""
    runs: list[tuple[str, bool]] = []
    bold = False
    for i, part in enumerate(s.split("**")):
        if part:
            runs.append((part, bold))
        bold = not bold if i < len(s.split("**")) - 1 else bold
    return runs or [("", False)]


def draft_to_pdf(draft: ResumeDraft) -> bytes:
    """Render a structured ResumeDraft to the single-column template PDF."""
    try:
        from fpdf import FPDF  # type: ignore
    except ImportError as e:
        raise RendererUnavailable("fpdf2 not installed; pip install fpdf2") from e

    pdf = FPDF(format="A4")
    pdf.set_auto_page_break(auto=True, margin=14)
    pdf.set_margins(16, 14, 16)
    pdf.add_page()
    epw = pdf.w - pdf.l_margin - pdf.r_margin

    # Header: centered name + contact line.
    pdf.set_text_color(15, 15, 22)
    pdf.set_font("Helvetica", "B", 22)
    pdf.cell(0, 10, _latin1(draft.candidate_name), align="C",
             new_x="LMARGIN", new_y="NEXT")
    contact = draft.contact_line()
    if contact:
        pdf.set_font("Helvetica", size=9.5)
        pdf.set_text_color(90, 92, 110)
        pdf.cell(0, 5, _latin1(contact), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    if draft.summary:
        pdf.set_text_color(40, 40, 48)
        pdf.set_font("Helvetica", size=9.5)
        pdf.multi_cell(0, 4.6, _latin1(draft.summary), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(1)

    def section_header(title: str) -> None:
        pdf.ln(2)
        pdf.set_font("Helvetica", "B", 12)
        pdf.set_text_color(15, 15, 22)
        pdf.cell(0, 6, _latin1(title), new_x="LMARGIN", new_y="NEXT")
        y = pdf.get_y()
        pdf.set_draw_color(40, 40, 48)
        pdf.set_line_width(0.4)
        pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
        pdf.ln(1.5)

    def row(left_md: str, right: str, size: float = 10.0) -> None:
        y = pdf.get_y()
        pdf.set_font("Helvetica", size=size)
        rw = (pdf.get_string_width(_latin1(right)) + 1) if right else 0.0
        max_left = epw - rw - 2
        # Render left runs inline (bold/regular), clipped to max_left.
        pdf.set_xy(pdf.l_margin, y)
        used = 0.0
        for text, bold in _parse_runs(left_md):
            pdf.set_font("Helvetica", "B" if bold else "", size)
            t = _latin1(text)
            w = pdf.get_string_width(t)
            if used + w > max_left:
                # truncate this run to fit
                while t and used + pdf.get_string_width(t) > max_left:
                    t = t[:-1]
                w = pdf.get_string_width(t)
            if t:
                pdf.cell(w, 5.4, t, new_x="RIGHT", new_y="TOP")
                used += w
            if used >= max_left:
                break
        if right:
            pdf.set_font("Helvetica", size=size)
            pdf.set_text_color(90, 92, 110)
            pdf.set_xy(pdf.w - pdf.r_margin - rw, y)
            pdf.cell(rw, 5.4, _latin1(right), align="R",
                     new_x="LMARGIN", new_y="NEXT")
            pdf.set_text_color(40, 40, 48)
        else:
            pdf.set_xy(pdf.l_margin, y + 5.4)

    for sec in draft.sections:
        if sec.kind == "summary":
            continue
        section_header(sec.title)
        pdf.set_text_color(40, 40, 48)
        if sec.kind == "skills" and sec.body:
            pdf.set_font("Helvetica", size=10)
            pdf.multi_cell(0, 5, _latin1(sec.body), new_x="LMARGIN", new_y="NEXT")
            continue
        if sec.body:
            pdf.set_font("Helvetica", size=10)
            pdf.multi_cell(0, 5, _latin1(sec.body), new_x="LMARGIN", new_y="NEXT")
        for r in sec.rows:
            right = str(r.get("right") or r.get("link") or "")
            row(str(r.get("left", "")), right)
            pdf.set_font("Helvetica", size=9.5)
            pdf.set_text_color(45, 45, 55)
            for b in r.get("bullets", []):
                text = b["text"] if isinstance(b, dict) else str(b)
                pdf.set_x(pdf.l_margin + 3)
                pdf.multi_cell(epw - 3, 4.6, _latin1("•  " + text),
                               new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1.2)

    return bytes(pdf.output())


def draft_to_styled_html(draft: ResumeDraft, footer: str | None = None) -> str:
    """Render a ResumeDraft to a self-contained single-column HTML page.

    ``footer``, when set, renders as a small centered line at the bottom of
    the page (e.g. a "Built with JobHunt" attribution on public share links).
    """
    def runs_to_html(s: str) -> str:
        out = []
        for text, bold in _parse_runs(s):
            t = _esc(text)
            out.append(f"<strong>{t}</strong>" if bold else t)
        return "".join(out)

    blocks: list[str] = []
    for sec in draft.sections:
        if sec.kind == "summary":
            continue
        blocks.append(f"<h2>{_esc(sec.title)}</h2>")
        if sec.kind == "skills" and sec.body:
            blocks.append(f'<p class="skills">{_esc(sec.body)}</p>')
            continue
        if sec.body:
            blocks.append(f"<p>{_esc(sec.body)}</p>")
        for r in sec.rows:
            right = str(r.get("right") or "")
            link = str(r.get("link") or "")
            right_html = (f'<a href="{_esc(link)}">{_esc(right or link)}</a>'
                          if link else _esc(right))
            blocks.append(
                '<div class="row"><span class="left">'
                f'{runs_to_html(str(r.get("left", "")))}</span>'
                f'<span class="right">{right_html}</span></div>'
            )
            bullets = r.get("bullets", [])
            if bullets:
                lis = "".join(
                    f"<li>{_esc(b['text'] if isinstance(b, dict) else str(b))}</li>"
                    for b in bullets
                )
                blocks.append(f"<ul>{lis}</ul>")

    contact = _esc(draft.contact_line())
    summary = f'<p class="summary">{_esc(draft.summary)}</p>' if draft.summary else ""
    inner = "\n  ".join(blocks)
    footer_html = f'<div class="footer">{_esc(footer)}</div>' if footer else ""
    return f"""<!doctype html>
<html><head><meta charset="utf-8">
<title>{_esc(draft.candidate_name)} — {_esc(draft.target_role)}</title>
<style>
  :root {{ --ink:#14161f; --muted:#5b6478; --rule:#14161f; }}
  body {{ font-family: "Helvetica Neue", Helvetica, Arial, sans-serif;
          max-width: 740px; margin: 40px auto; padding: 0 28px;
          color: var(--ink); line-height: 1.45; }}
  h1 {{ text-align:center; font-size: 2.0em; margin: 0 0 2px; letter-spacing:-0.01em; }}
  .contact {{ text-align:center; color: var(--muted); font-size: 0.86em; margin-bottom: 14px; }}
  .summary {{ font-size: 0.92em; margin: 0 0 6px; }}
  h2 {{ font-size: 1.05em; margin: 18px 0 2px; padding-bottom: 3px;
        border-bottom: 1.5px solid var(--rule); }}
  .row {{ display:flex; justify-content:space-between; gap:12px;
          margin-top: 8px; font-size: 0.92em; }}
  .row .left strong {{ font-weight: 700; }}
  .row .right {{ color: var(--muted); white-space: nowrap; font-size: 0.92em; }}
  .row .right a {{ color: var(--muted); text-decoration: none; }}
  ul {{ margin: 3px 0 4px; padding-left: 18px; }}
  li {{ margin-bottom: 3px; font-size: 0.9em; }}
  p.skills {{ font-size: 0.9em; margin: 4px 0; }}
  .footer {{ text-align:center; color: var(--muted); font-size: 0.78em;
             margin-top: 28px; padding-top: 10px; border-top: 1px solid #e3e6ef; }}
</style></head><body>
  <h1>{_esc(draft.candidate_name)}</h1>
  <div class="contact">{contact}</div>
  {summary}
  {inner}
  {footer_html}
</body></html>"""


def draft_to_docx(draft: ResumeDraft) -> bytes:
    """Render a ResumeDraft to a .docx matching the single-column template."""
    try:
        from docx import Document  # type: ignore
        from docx.enum.text import WD_TAB_ALIGNMENT
        from docx.shared import Inches, Pt
    except ImportError as e:
        raise RendererUnavailable(
            "python-docx not installed; pip install python-docx"
        ) from e

    doc = Document()
    name = doc.add_paragraph()
    name.alignment = 1  # center
    run = name.add_run(draft.candidate_name)
    run.bold = True
    run.font.size = Pt(20)
    contact = draft.contact_line()
    if contact:
        c = doc.add_paragraph()
        c.alignment = 1
        c.add_run(contact).font.size = Pt(9)
    if draft.summary:
        doc.add_paragraph(draft.summary)

    right_tab = Inches(6.5)
    for sec in draft.sections:
        if sec.kind == "summary":
            continue
        h = doc.add_heading(sec.title, level=2)
        h.paragraph_format.space_before = Pt(8)
        if sec.kind == "skills" and sec.body:
            doc.add_paragraph(sec.body)
            continue
        if sec.body:
            doc.add_paragraph(sec.body)
        for r in sec.rows:
            p = doc.add_paragraph()
            p.paragraph_format.tab_stops.add_tab_stop(
                right_tab, WD_TAB_ALIGNMENT.RIGHT)
            for text, bold in _parse_runs(str(r.get("left", ""))):
                run = p.add_run(text)
                run.bold = bold
            right = str(r.get("right") or r.get("link") or "")
            if right:
                p.add_run("\t" + right)
            for b in r.get("bullets", []):
                text = b["text"] if isinstance(b, dict) else str(b)
                doc.add_paragraph(text, style="List Bullet")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------- factory


def get_renderer(fmt: str) -> Renderer:
    """Return a renderer by extension. Falls back to text on unknown fmt."""
    table: dict[str, type] = {
        "txt": TextRenderer,
        "html": HTMLRenderer,
        "pdf": PDFRenderer,
        "docx": DOCXRenderer,
    }
    cls = table.get(fmt.lower(), TextRenderer)
    return cls()
